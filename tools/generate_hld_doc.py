from pathlib import Path
from datetime import date
import math

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\ACER\manufacturing-infrastructure-chatbot")
DOCS = ROOT / "docs"
ASSETS = DOCS / "hld_assets"
OUT = DOCS / "InfraChat_HLD_IBM_Format.docx"
CSV = ROOT / "data" / "predictive_maintenance.csv"


def load_font(name, size):
    try:
        return ImageFont.truetype(name, size)
    except Exception:
        return ImageFont.load_default()


FONT = load_font("arial.ttf", 24)
FONT_B = load_font("arialbd.ttf", 26)
FONT_S = load_font("arial.ttf", 18)

COLORS = {
    "bg": "#F8FAFC",
    "line": "#334155",
    "sky": "#E0F2FE",
    "green": "#ECFDF5",
    "amber": "#FEF3C7",
    "card": "#FFFFFF",
    "text": "#0F172A",
    "red": "#EF4444",
}


def draw_box(draw, xy, text, fill="#FFFFFF", outline="#334155"):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = (current + " " + word).strip()
        if draw.textlength(test, font=FONT_B) > (x2 - x1 - 30) and current:
            lines.append(current)
            current = word
        else:
            current = test
    if current:
        lines.append(current)
    y = y1 + ((y2 - y1) - len(lines) * 30) // 2
    for line in lines:
        width = draw.textlength(line, font=FONT_B)
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=FONT_B, fill=COLORS["text"])
        y += 30


def arrow(draw, start, end):
    draw.line([start, end], fill=COLORS["line"], width=4)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    for delta in (2.65, -2.65):
        draw.line(
            [end, (x2 + 18 * math.cos(angle + delta), y2 + 18 * math.sin(angle + delta))],
            fill=COLORS["line"],
            width=4,
        )


def make_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)

    img = Image.new("RGB", (1400, 650), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((40, 35), "System Context Diagram", font=FONT_B, fill=COLORS["text"])
    boxes = [
        ((70, 240, 310, 360), "User / Operator", COLORS["sky"]),
        ((410, 240, 650, 360), "Web Browser", COLORS["green"]),
        ((750, 210, 1040, 390), "InfraChat Flask Application", COLORS["card"]),
        ((1140, 240, 1340, 360), "CSV Dataset", COLORS["amber"]),
    ]
    for xy, text, color in boxes:
        draw_box(d, xy, text, color)
    arrow(d, (310, 300), (410, 300))
    arrow(d, (650, 300), (750, 300))
    arrow(d, (1040, 300), (1140, 300))
    img.save(ASSETS / "system_context.png")

    img = Image.new("RGB", (1400, 820), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((40, 35), "High Level Architecture", font=FONT_B, fill=COLORS["text"])
    layers = [
        ((60, 110, 1340, 280), "Presentation Layer", COLORS["sky"], ["Home Page", "Dashboard Page", "script.js", "dashboard.js", "Chart.js"]),
        ((60, 330, 1340, 520), "Application Layer", COLORS["green"], ["Flask Routes", "Response Engine", "Fuzzy Matcher", "Risk Calculator", "Dashboard API"]),
        ((60, 570, 1340, 740), "Data Layer", COLORS["amber"], ["predictive_maintenance.csv", "Pandas DataFrame", "In-memory Session Context"]),
    ]
    for rect, title, color, items in layers:
        d.rounded_rectangle(rect, radius=20, fill=color, outline=COLORS["line"], width=2)
        d.text((rect[0] + 25, rect[1] + 18), title, font=FONT_B, fill=COLORS["text"])
        x = rect[0] + 45
        y = rect[1] + 75
        for item in items:
            draw_box(d, (x, y, x + 220, y + 70), item, COLORS["card"])
            x += 245
    arrow(d, (700, 280), (700, 330))
    arrow(d, (700, 520), (700, 570))
    img.save(ASSETS / "architecture.png")

    img = Image.new("RGB", (1400, 820), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((40, 35), "Chat Query Process Flow", font=FONT_B, fill=COLORS["text"])
    steps = [
        ((80, 130, 330, 230), "User enters query", COLORS["sky"]),
        ((430, 130, 680, 230), "POST /chat", COLORS["card"]),
        ((780, 130, 1130, 230), "Normalize text and scan ID", COLORS["card"]),
        ((780, 330, 1130, 430), "Product ID found?", COLORS["amber"]),
        ((430, 330, 680, 430), "Fetch machine record", COLORS["green"]),
        ((80, 330, 330, 430), "Return machine details", COLORS["green"]),
        ((780, 550, 1130, 650), "Fuzzy command match", COLORS["card"]),
        ((430, 550, 680, 650), "Calculate aggregate / metric", COLORS["card"]),
        ((80, 550, 330, 650), "Return chatbot answer", COLORS["sky"]),
    ]
    for xy, text, color in steps:
        draw_box(d, xy, text, color)
    for start, end in [
        ((330, 180), (430, 180)),
        ((680, 180), (780, 180)),
        ((955, 230), (955, 330)),
        ((780, 380), (680, 380)),
        ((430, 380), (330, 380)),
        ((955, 430), (955, 550)),
        ((780, 600), (680, 600)),
        ((430, 600), (330, 600)),
    ]:
        arrow(d, start, end)
    d.text((1000, 470), "No", font=FONT_S, fill=COLORS["red"])
    d.text((710, 350), "Yes", font=FONT_S, fill="#059669")
    img.save(ASSETS / "chat_process_flow.png")

    img = Image.new("RGB", (1400, 720), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((40, 35), "Information Flow", font=FONT_B, fill=COLORS["text"])
    boxes = [
        ((80, 160, 350, 270), "CSV machine records", COLORS["amber"]),
        ((480, 120, 790, 220), "Pandas processing", COLORS["card"]),
        ((480, 300, 790, 400), "Chat response logic", COLORS["green"]),
        ((920, 120, 1240, 220), "Dashboard JSON", COLORS["sky"]),
        ((920, 300, 1240, 400), "Chatbot response JSON", COLORS["sky"]),
        ((500, 520, 780, 620), "Browser UI renders output", COLORS["card"]),
    ]
    for xy, text, color in boxes:
        draw_box(d, xy, text, color)
    for start, end in [
        ((350, 215), (480, 170)),
        ((350, 215), (480, 350)),
        ((790, 170), (920, 170)),
        ((790, 350), (920, 350)),
        ((1080, 220), (700, 520)),
        ((1080, 400), (700, 520)),
    ]:
        arrow(d, start, end)
    img.save(ASSETS / "information_flow.png")

    img = Image.new("RGB", (1400, 650), COLORS["bg"])
    d = ImageDraw.Draw(img)
    d.text((40, 35), "Deployment View", font=FONT_B, fill=COLORS["text"])
    d.rounded_rectangle((70, 120, 1330, 560), radius=22, fill="#FFFFFF", outline=COLORS["line"], width=3)
    d.text((100, 150), "Local Development / Demo Host", font=FONT_B, fill=COLORS["text"])
    boxes = [
        ((130, 260, 390, 380), "Browser Client", COLORS["sky"]),
        ((540, 230, 860, 410), "Python Flask Server app.py", COLORS["green"]),
        ((1020, 260, 1250, 380), "Local CSV File", COLORS["amber"]),
    ]
    for xy, text, color in boxes:
        draw_box(d, xy, text, color)
    arrow(d, (390, 320), (540, 320))
    arrow(d, (860, 320), (1020, 320))
    d.text((430, 285), "HTTP routes: /, /dashboard, /chat, /api/dashboard-data", font=FONT_S, fill=COLORS["line"])
    img.save(ASSETS / "deployment_view.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True


def para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph("")
    return table


def add_pic(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


def add_title_page(doc):
    for _ in range(3):
        doc.add_paragraph("")
    for text, size, bold in [
        ("University Name", 18, True),
        ("Project - High Level Design", 22, True),
        ("on", 11, False),
        ("Manufacturing Infrastructure Chatbot for Predictive Maintenance Monitoring", 18, True),
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(size)
        p.runs[0].bold = bold
    for _ in range(3):
        doc.add_paragraph("")
    p = doc.add_paragraph("Student Name: ______________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(["Industry Mentor:", "", "Guided by:", ""]):
        table.cell(0, i).text = value
    doc.add_paragraph("")
    p = doc.add_paragraph(f"Document Version: 1.0    Date: {date.today().strftime('%d-%m-%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_toc(doc):
    doc.add_paragraph("Table of Contents", style="Heading 1")
    items = [
        "Introduction",
        "Scope of the document",
        "Intended Audience",
        "System overview",
        "System Design",
        "Application Design",
        "Process Flow",
        "Information Flow",
        "Components Design",
        "Key Design Considerations",
        "API Catalogue",
        "Data Design",
        "Data Model",
        "Data Access Mechanism",
        "Data Retention Policies",
        "Data Migration",
        "Interfaces",
        "State and Session Management",
        "Caching",
        "Non-Functional Requirements",
        "Security Aspects",
        "Performance Aspects",
        "Recommended Pictures and Diagrams",
        "References",
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Paragraph")
    doc.add_page_break()


def build_doc():
    df = pd.read_csv(CSV)
    total = len(df)
    failed = int(df["Target"].sum())
    working = total - failed
    failure_rate = round((failed / total) * 100, 2)
    avg_air = round(df["Air temperature [K]"].mean(), 1)
    avg_process = round(df["Process temperature [K]"].mean(), 1)
    avg_rpm = int(round(df["Rotational speed [rpm]"].mean(), 0))
    avg_torque = round(df["Torque [Nm]"].mean(), 1)
    avg_wear = round(df["Tool wear [min]"].mean(), 1)
    high_wear = int((df["Tool wear [min]"] > 200).sum())

    doc = Document()
    style_doc(doc)
    doc.sections[0].footer.paragraphs[0].text = "Manufacturing Infrastructure Chatbot - High Level Design"
    add_title_page(doc)
    add_toc(doc)

    opening = [
        ("1. Introduction", "This High Level Design document describes the architecture, major modules, information flow, interfaces, data design, and non-functional design considerations for the Manufacturing Infrastructure Chatbot, also referred to as InfraChat. The application is a web-based machine monitoring assistant that allows users to ask natural-language questions about manufacturing machine health and view analytics through a dashboard."),
        ("2. Scope of the document", "This document covers the high-level design of the current project implementation, including Flask routes, frontend screens, chatbot query handling, dashboard analytics, CSV-based data access, and deployment assumptions. It does not cover production cloud deployment, enterprise authentication, real-time IoT ingestion, or machine learning model training because those capabilities are outside the current codebase."),
        ("3. Intended Audience", "The intended audience includes project reviewers, academic evaluators, developers, system designers, mentors, and stakeholders who need to understand the architecture and design choices of the InfraChat project."),
        ("4. System overview", f"InfraChat provides two main user experiences: a chatbot for plain-English maintenance queries and a dashboard for visual monitoring. The backend loads {total:,} predictive maintenance records from a CSV file and uses pandas for analytics. The chatbot uses rule-based intent handling and RapidFuzz-based fuzzy matching to answer questions about machine counts, failures, averages, machine IDs, and risk status. The dashboard exposes aggregated JSON data for KPI cards, charts, gauges, and a machine table."),
    ]
    for heading, text in opening:
        doc.add_paragraph(heading, style="Heading 1")
        para(doc, text)

    add_table(doc, ["Metric", "Current Value"], [
        ["Total machines", f"{total:,}"],
        ["Working machines", f"{working:,}"],
        ["Failed machines", f"{failed:,}"],
        ["Failure rate", f"{failure_rate}%"],
        ["Average air temperature", f"{avg_air} K"],
        ["Average process temperature", f"{avg_process} K"],
        ["Average RPM", avg_rpm],
        ["Average torque", f"{avg_torque} Nm"],
        ["Average tool wear", f"{avg_wear} min"],
        ["Machines with tool wear above 200 min", high_wear],
    ])

    doc.add_paragraph("5. System Design", style="Heading 1")
    para(doc, "The system follows a simple monolithic web application design. The browser sends page and API requests to the Flask backend. The backend performs command matching, Product ID lookup, risk calculation, and dashboard aggregation using an in-memory pandas DataFrame loaded from the CSV dataset.")
    add_pic(doc, ASSETS / "system_context.png", "Figure 1: System context diagram")
    add_pic(doc, ASSETS / "architecture.png", "Figure 2: High level architecture diagram")

    doc.add_paragraph("6. Application Design", style="Heading 1")
    add_table(doc, ["Layer", "Project Files / Components", "Responsibility"], [
        ["Presentation Layer", "templates/index.html, templates/dashboard.html, static/style.css", "Provides home page, dashboard page, navigation, chatbot window, and visual layout."],
        ["Client Logic", "static/script.js, static/dashboard.js", "Sends chat requests, loads dashboard JSON, renders charts, updates KPI cards, and builds machine table."],
        ["Application Layer", "app.py", "Defines Flask routes, loads data, processes queries, calculates risk, and builds API responses."],
        ["Data Layer", "data/predictive_maintenance.csv", "Stores machine records, sensor values, target status, and failure type information."],
        ["External Libraries", "Flask, pandas, RapidFuzz, Chart.js", "Provide web routing, data processing, fuzzy matching, and chart rendering."],
    ])

    doc.add_paragraph("7. Process Flow", style="Heading 1")
    para(doc, "The main process begins when a user sends a message through the chatbot. The browser posts the message to the Flask /chat route. The backend normalizes the input, searches for a Product ID, applies fuzzy command matching when needed, reads the required values from the DataFrame, and returns a JSON response for display.")
    add_pic(doc, ASSETS / "chat_process_flow.png", "Figure 3: Chat query process flow")

    doc.add_paragraph("8. Information Flow", style="Heading 1")
    para(doc, "Machine data flows from the CSV file into an in-memory pandas DataFrame at application startup. Chatbot responses and dashboard analytics both use this shared DataFrame. Dashboard requests return structured JSON, while chatbot requests return formatted response text.")
    add_pic(doc, ASSETS / "information_flow.png", "Figure 4: Information flow across data, backend, and browser")

    doc.add_paragraph("9. Components Design", style="Heading 1")
    add_table(doc, ["Component", "Description", "Key Design Notes"], [
        ["Flask Web Server", "Hosts pages and APIs.", "Runs as a single Python application in the current project."],
        ["Chat Response Engine", "Handles user queries and returns text responses.", "Uses Product ID detection, fuzzy command matching, and context-aware follow-up handling."],
        ["Command Dictionary", "Maps intents to supported phrases.", "Improves usability for non-technical query wording."],
        ["Risk Calculator", "Classifies machine condition.", "Uses Target, tool wear, air temperature, and torque thresholds."],
        ["Dashboard API", "Builds KPI and chart payloads.", "Returns JSON for machine status, failure types, machine types, tool wear, averages, and sample records."],
        ["Frontend Dashboard", "Renders visual analytics.", "Uses Chart.js doughnut, bar, and gauge-like visualizations."],
        ["CSV Data Store", "Stores machine records.", "Loaded once at startup and reused for all requests."],
    ])

    doc.add_paragraph("10. Key Design Considerations", style="Heading 1")
    bullets(doc, [
        "The system is intentionally lightweight and suitable for academic demonstration or local use.",
        "CSV storage keeps setup simple, but a database should be used for production scale.",
        "Fuzzy matching improves user experience without requiring a full NLP or generative AI model.",
        "The current in-memory session context supports simple follow-up queries but is not multi-user safe.",
        "Dashboard calculations are fast because the dataset is loaded into memory and has limited size.",
        "The application should be hardened with authentication, logging, exception handling, and HTTPS before production use.",
    ])

    doc.add_paragraph("11. API Catalogue", style="Heading 1")
    add_table(doc, ["API / Route", "Method", "Purpose", "Response"], [
        ["/", "GET", "Loads the home page with chatbot entry point.", "HTML page"],
        ["/dashboard", "GET", "Loads analytics dashboard page.", "HTML page"],
        ["/api/dashboard-data", "GET", "Returns KPI, chart, table, and at-risk machine data.", "JSON payload"],
        ["/chat", "POST", "Accepts a user message and returns chatbot response.", "JSON with response field"],
    ])
    para(doc, 'Example /chat request: { "message": "failed machines" }. Example /chat response: { "response": "Total failed machines: 339" }.')

    doc.add_paragraph("12. Data Design", style="Heading 1")
    para(doc, "The project uses a predictive maintenance CSV dataset as its primary data source. Each record represents a machine/product with sensor values and failure classification fields.")
    add_table(doc, ["Column", "Purpose"], [
        ["UDI", "Unique row identifier."],
        ["Product ID", "Machine or product identifier used for lookup."],
        ["Type", "Machine quality type: L, M, or H."],
        ["Air temperature [K]", "Air temperature reading."],
        ["Process temperature [K]", "Process temperature reading."],
        ["Rotational speed [rpm]", "Machine rotational speed."],
        ["Torque [Nm]", "Torque/load value."],
        ["Tool wear [min]", "Tool usage duration."],
        ["Target", "Failure indicator: 1 means failed, 0 means working."],
        ["Failure Type", "Specific failure category or No Failure."],
    ])

    doc.add_paragraph("13. Data Model", style="Heading 1")
    add_table(doc, ["Entity", "Attributes", "Relationship"], [
        ["MachineRecord", "UDI, Product ID, Type, temperatures, RPM, torque, tool wear, Target, Failure Type", "One row represents one machine observation used by chatbot and dashboard."],
        ["DashboardSummary", "total, working, failed, failure_rate, averages, chart labels, chart values", "Derived from MachineRecord records at request time."],
        ["SessionContext", "last_machine_id", "Stores the latest referenced Product ID for follow-up chat queries."],
    ])

    for heading, text in [
        ("14. Data Access Mechanism", "The dataset is loaded using pandas.read_csv during application startup. All filtering and aggregation is performed through pandas DataFrame operations such as boolean filters, mean(), sum(), value_counts(), sort_values(), and to_dict()."),
        ("15. Data Retention Policies", "The current application does not create, update, or delete dataset records. The CSV file is retained in the project data folder and remains unchanged during normal application use. Chat session context is temporary and exists only in application memory until the process stops."),
        ("16. Data Migration", "No automated migration is required for the current CSV-based design. If the solution is upgraded to a relational database, the CSV schema can be mapped to a MachineRecord table and loaded using a one-time import script. Future migrations should preserve Product ID, sensor values, target status, and failure type fields."),
    ]:
        doc.add_paragraph(heading, style="Heading 1")
        para(doc, text)

    doc.add_paragraph("17. Interfaces", style="Heading 1")
    add_table(doc, ["Interface", "Consumer", "Provider", "Description"], [
        ["Browser UI", "User", "HTML/CSS/JavaScript", "Allows user navigation, chatbot input, dashboard viewing, and data refresh."],
        ["Chat API", "script.js", "Flask /chat route", "Sends user message and receives response text."],
        ["Dashboard API", "dashboard.js", "Flask /api/dashboard-data route", "Receives analytics payload for charts and tables."],
        ["Data Interface", "Flask app", "CSV file through pandas", "Reads machine records for lookup and analytics."],
    ])

    for heading, text in [
        ("18. State and Session Management", "The application stores minimal state in SESSION_CONTEXT with last_machine_id. This supports follow-up questions such as rpm, torque, temperature, or risk after a user has searched for a Product ID. For production, this should be replaced with per-user Flask sessions or database-backed session storage."),
        ("19. Caching", "The CSV dataset is effectively cached in memory because it is loaded into a pandas DataFrame once when app.py starts. No Redis, browser cache strategy, or external caching layer is currently implemented."),
    ]:
        doc.add_paragraph(heading, style="Heading 1")
        para(doc, text)

    doc.add_paragraph("20. Non-Functional Requirements", style="Heading 1")
    add_table(doc, ["Area", "Requirement / Consideration"], [
        ["Usability", "Users should be able to ask simple natural-language queries and view dashboard charts without technical skills."],
        ["Reliability", "The application depends on the Flask server and availability of the CSV file."],
        ["Maintainability", "The current single-file backend is easy to inspect; future growth should split routes, services, and utilities."],
        ["Scalability", "Suitable for local/demo datasets; production scale should use a database and possibly background analytics jobs."],
        ["Compatibility", "Frontend is browser-based and backend is Python/Flask based."],
    ])

    doc.add_paragraph("21. Security Aspects", style="Heading 1")
    bullets(doc, [
        "The current project assumes trusted local/demo usage.",
        "No login, authorization, HTTPS enforcement, CSRF protection, or API rate limiting is implemented.",
        "User input is processed as text and should be validated consistently before production release.",
        "Production deployment should add authentication, role-based access, secure headers, logging, and protected configuration management.",
    ])

    doc.add_paragraph("22. Performance Aspects", style="Heading 1")
    bullets(doc, [
        "Dataset is loaded once at startup, avoiding repeated disk reads for every query.",
        "Aggregations over the current 10,000-row dataset are lightweight.",
        "RapidFuzz command matching is acceptable for the current command dictionary size.",
        "Dashboard chart data is calculated on request; this is fine for demo use but can be cached for larger datasets.",
    ])

    doc.add_paragraph("23. Recommended Pictures and Diagrams", style="Heading 1")
    para(doc, "Use the following pictures and diagrams in the final submission. The first five diagrams have already been added to this document. Screenshots should be captured from your running application for the most accurate project evidence.")
    add_table(doc, ["Picture / Diagram", "Where to Use", "What It Should Show"], [
        ["System Context Diagram", "System Design section", "User, browser, Flask application, and CSV dataset."],
        ["High Level Architecture Diagram", "System Design section", "Presentation layer, application layer, and data layer."],
        ["Chat Query Process Flow", "Process Flow section", "Input query, /chat API, Product ID detection, fuzzy matching, and response."],
        ["Information Flow Diagram", "Information Flow section", "CSV to pandas DataFrame to chatbot/dashboard JSON to browser UI."],
        ["Deployment View", "System Design or References section", "Browser, Flask server, and local CSV file running on one host."],
        ["Home Page Screenshot", "System overview or Application Design", "InfraChat landing page with sidebar and chatbot button."],
        ["Chatbot Screenshot", "Application Design or Interfaces", "Chat window showing sample questions and a machine response."],
        ["Dashboard Screenshot", "Application Design or Information Flow", "KPI cards, RPM/torque gauges, failure chart, and machine table."],
        ["Dataset Sample Screenshot", "Data Design section", "CSV columns including Product ID, sensor values, Target, and Failure Type."],
    ])
    add_pic(doc, ASSETS / "deployment_view.png", "Figure 5: Current deployment view")

    doc.add_paragraph("24. References", style="Heading 1")
    bullets(doc, [
        str(ROOT / "app.py"),
        str(ROOT / "templates" / "index.html"),
        str(ROOT / "templates" / "dashboard.html"),
        str(ROOT / "static" / "script.js"),
        str(ROOT / "static" / "dashboard.js"),
        str(ROOT / "data" / "predictive_maintenance.csv"),
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_diagrams()
    build_doc()
